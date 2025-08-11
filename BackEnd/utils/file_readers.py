# utils/file_readers.py
import io
import os
import docx
import pandas as pd
from PyPDF2 import PdfReader
from fastapi import UploadFile
from typing import List
from langchain.docstore.document import Document

# OCR optional
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False


async def extract_documents(files: List[UploadFile]) -> List[Document]:
    """
    Reads files and returns list of langchain Documents with metadata:
      - source: original filename
      - basename: filename without extension (lowercase)
      - doc_id: string index (1-based) as sent in upload order
    """
    documents: List[Document] = []
    for idx, file in enumerate(files):
        content = await file.read()
        filename = file.filename or f"file_{idx+1}"
        lower = filename.lower()
        text = ""

        try:
            if lower.endswith(".pdf"):
                pdf_stream = io.BytesIO(content)
                reader = PdfReader(pdf_stream)
                text = "\n".join([p.extract_text() or "" for p in reader.pages])

            elif lower.endswith((".docx", ".doc")):
                doc_stream = io.BytesIO(content)
                doc = docx.Document(doc_stream)
                text = "\n".join([p.text for p in doc.paragraphs if p.text])

            elif lower.endswith(".xlsx"):
                df_stream = io.BytesIO(content)
                df = pd.read_excel(df_stream)
                text = df.astype(str).apply(lambda x: " ".join(x), axis=1).str.cat(sep="\n")

            elif lower.endswith(".csv"):
                csv_stream = io.BytesIO(content)
                df = pd.read_csv(csv_stream)
                text = df.astype(str).apply(lambda x: " ".join(x), axis=1).str.cat(sep="\n")

            elif lower.endswith((".png", ".jpg", ".jpeg", ".tiff")):
                if OCR_AVAILABLE:
                    img_stream = io.BytesIO(content)
                    img = Image.open(img_stream)
                    text = pytesseract.image_to_string(img)
                else:
                    text = "[Image uploaded — OCR not available]"

            else:
                text = content.decode("utf-8", errors="ignore")

        except Exception as e:
            print(f"[file_readers] Error processing {filename}: {e}")
            text = f"[Error reading file: {filename}]"

        if text:
            basename = os.path.splitext(filename)[0].strip().lower()
            metadata = {"source": filename, "basename": basename, "doc_id": str(idx + 1)}
            documents.append(Document(page_content=text, metadata=metadata))

    return documents
